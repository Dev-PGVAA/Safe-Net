from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/artem/Documents/0. Projects/2. study/safe-net")
OUT = ROOT / "deliverables/Документация проекта Safe-Net — конференция.docx"
PHOTO = ROOT / "artifacts/safe-net-doc-review/cyber-learning-photo.png"
DIAGRAM = ROOT / "artifacts/safe-net-doc-review/architecture.png"


def font(size, bold=False):
    face = "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"
    return ImageFont.truetype(face, size)


def rounded(draw, box, fill, outline=None, radius=28):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)


def center(draw, box, text, fill, fnt):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, align="center", spacing=8)
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=fnt, fill=fill, align="center", spacing=8)


def create_diagram():
    image = Image.new("RGB", (1600, 900), "#f8f7ff")
    draw = ImageDraw.Draw(image)
    title = "Архитектура SafeNet"
    draw.text((80, 56), title, font=font(54, True), fill="#29213d")
    draw.text((82, 128), "Единый модуль правил соединяет обучение, симулятор и защиту ссылок.", font=font(27), fill="#675d80")
    boxes = [
        (90, 285, 430, 475, "Web-клиент\nNext.js + React", "#e9e4ff"),
        (630, 285, 970, 475, "API\nNestJS", "#dcefff"),
        (1170, 285, 1510, 475, "База данных\nPostgreSQL", "#dcf7ec"),
        (360, 640, 760, 815, "@safe-net/guard-core\nединые правила проверки", "#ffe8ba"),
        (930, 640, 1320, 815, "Guard и ML-слой\nдополнительная защита", "#ffe0ed"),
    ]
    for left, top, right, bottom, label, color in boxes:
        rounded(draw, (left, top, right, bottom), color, "#ddd6eb")
        center(draw, (left, top, right, bottom), label, "#29213d", font(31, True))
    arrows = [((430, 380), (630, 380)), ((970, 380), (1170, 380)), ((530, 475), (550, 640)), ((1050, 475), (1125, 640)), ((760, 727), (930, 727))]
    for start, end in arrows:
        draw.line([start, end], fill="#7c5cff", width=10)
        dx, dy = end[0] - start[0], end[1] - start[1]
        if abs(dx) >= abs(dy):
            sign = 1 if dx > 0 else -1
            draw.polygon([(end[0], end[1]), (end[0] - sign * 22, end[1] - 14), (end[0] - sign * 22, end[1] + 14)], fill="#7c5cff")
        else:
            sign = 1 if dy > 0 else -1
            draw.polygon([(end[0], end[1]), (end[0] - 14, end[1] - sign * 22), (end[0] + 14, end[1] - sign * 22)], fill="#7c5cff")
    image.save(DIAGRAM)


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    props.append(shd)


def move_after(anchor, block):
    anchor._p.addnext(block._element)


def add_caption(document, anchor, text):
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.left_indent = Cm(0)
    caption.paragraph_format.right_indent = Cm(0)
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    move_after(anchor, caption)
    return caption


def add_image(document, anchor, path, caption, width_cm):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    move_after(anchor, paragraph)
    add_caption(document, paragraph, caption)
    return paragraph


def add_placeholder(document, anchor, title, instruction):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, "EEEAFE")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(16)
    heading = paragraph.add_run(title + "\n")
    heading.bold = True
    heading.font.size = Pt(13)
    heading.font.color.rgb = RGBColor(80, 58, 155)
    body = paragraph.add_run(instruction)
    body.font.size = Pt(10)
    move_after(anchor, table)
    return table


def replace(document, old, new):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return paragraph
    raise ValueError(f"Missing paragraph: {old}")


create_diagram()
document = Document(OUT)

# The source has a manually typed table of contents, so its page numbers need
# to follow the additional figures and visual-plan page.
toc_pages = {
    "1. Введение": 3,
    "2. Цели и задачи работы": 4,
    "3. Методика выполнения работы": 5,
    "4. Результаты работы и их проверка (испытания, апробация)": 7,
    "5. Выводы и перспективы дальнейшей работы": 8,
    "6. Список используемой литературы": 9,
}
for paragraph in document.paragraphs:
    if paragraph.style.name == "toc 1":
        raw = paragraph.text.strip()
        for label, page in toc_pages.items():
            if label in raw:
                number = label.split(".", 1)[0]
                title = label.split(". ", 1)[1]
                paragraph.text = f"{number}.\t{title}\t{page}"
                break
            title = label.split(". ", 1)[1]
            if title in raw:
                number = label.split(".", 1)[0]
                paragraph.text = f"{number}.\t{title}\t{page}"
                break

# Factual corrections aligned with the project architecture.
replace(document, "Создать backend на NextJS с REST API.", "Разработать серверную часть на NestJS с REST API.")
replace(document, "Создание REST API на NextJS (регистрация, логин, сохранение прогресса);", "Создание REST API на NestJS (регистрация, вход, сохранение прогресса);")
replace(document, "Server-Side Rendering (SSR) улучшает SEO и начальную загрузку страниц, что важно для образовательной платформы. Встроенный роутинг и API routes упрощают архитектуру. App Router с Server Components позволяет оптимизировать производительность, загружая данные на сервере.", "Next.js используется для клиентской части: он обеспечивает маршрутизацию, адаптивный интерфейс и быструю загрузку страниц. Серверная бизнес-логика и REST API вынесены в NestJS, поэтому роли компонентов разделены и приложение проще развивать.")

paragraphs = {paragraph.text.strip(): paragraph for paragraph in document.paragraphs if paragraph.text.strip()}
intro_anchor = paragraphs["SafeNet делает обучение понятным, доступным и интерактивным, позволяет в простой форме объяснить как защититься от кибермошенников."]
protocol_anchor = paragraphs["Проведение апробации среди учеников школы."]
results_anchor = paragraphs["85% участников отметили, что стали лучше распознавать фишинг и мошенничество."]

add_image(document, intro_anchor, PHOTO, "Рисунок 1 — Иллюстрация актуальности темы кибербезопасности (генеративное изображение).", 15.8)
diagram_anchor = add_image(document, protocol_anchor, DIAGRAM, "Рисунок 2 — Архитектура SafeNet: клиент, API, база данных и общий модуль защитных правил.", 16.0)

placeholder_one = add_placeholder(
    document,
    results_anchor,
    "МЕСТО ДЛЯ СКРИНШОТА 1 — ГЛАВНЫЙ ЭКРАН ТРЕНАЖЁРА",
    "Вставить реальный скрин главной страницы SafeNet: карточка задания «Уровень 1: Фишинг», кнопки «Безопасно» и «Опасно».\nПодпись: «Рисунок 3 — Интерактивное задание по распознаванию фишинга»."
)
placeholder_two = add_placeholder(
    document,
    results_anchor,
    "МЕСТО ДЛЯ СКРИНШОТА 2 — SAFENET GUARD",
    "Вставить реальный скрин проверки URL с объяснением подозрительного домена.\nПодпись: «Рисунок 4 — Локальная проверка ссылки и объяснение факторов риска»."
)
# Place the second placeholder after the first one rather than both directly after the same anchor.
placeholder_one._element.addnext(placeholder_two._element)

# Add a concise conference-oriented callout before the conclusion.
conclusion = paragraphs["Выводы и перспективы дальнейшей работы"]
callout = document.add_paragraph()
callout.paragraph_format.space_before = Pt(8)
callout.paragraph_format.space_after = Pt(10)
run = callout.add_run("Для выступления: показать два реальных скриншота, кратко продемонстрировать проверку подозрительной ссылки и объяснить, как единый модуль правил связывает урок, тренажёр и Guard.")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 58, 155)
move_after(conclusion, callout)

document.save(OUT)
print(OUT)
